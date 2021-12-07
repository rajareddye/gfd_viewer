"""
Author : Rajareddy Eddulannagari
Date: 12_Mar_2021

A style in Word is such a set of specifications that may be applied, all at once, to a document element. Word has
paragraph styles, 
character styles, 
table styles, and 
numbering definitions. These are applied to a paragraph, a span of
text, a table, and a list, respectively.

Description:
    This tool converts SRD docx to PSP file format for notepad++ (PSP : pseudocode file extension)
    this ensures hierarchy of pspecs, procedures and diagrams as per the SRD docx
	PSP files can be opened in Notepad++ with pseudocode language plugin (PSP) for pseudocode syntax highlighting"

Usage:
    Inputs:
        Rs_T_A400392-xx_accepted is SRD with one or all CRs markup selected, and used "Accept All changes shown" and saved
        Rs_T_A400392-xx_rejected is SRD with one or all CRs markup selected, and used "Reject All changes shown" and saved

    Commands:
        $python extract_all_pspecs.py Rs_T_A400392-xx_accepted.docx
        OR
        $python extract_all_pspecs.py Rs_T_A400392-xx_accepted.docx Rs_T_A400392-xx_rejected.docx

    Output:
        This tool will create "accepted" and / or "rejected" folders with all PSP files , based on number of docx files
        passed as arguments , as mentioned in Commands section above.
        Now these folders can be directly compared with beyond compare to see the impacted changes(select compare by content)
        You can add the pseudo code grammar for beyondcompare for better readability.

Python-docx info
    Each line in docx file is considered as paragraph in python-docx, this paragraph contains
     all styles related to that line of text.

"""
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import inspect
import os
import docx
import shutil
import sys
import time
import datetime
import re
from docx import Document
from docx.document import Document as _Document
from docx.image.image import Image
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


heading_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6", "Heading 7",
                  "Heading 8", "Heading 9", "Title1", "Title2", "Title3"]
comment_styles = ["Req_ID_note", "Req_Properties", "toc 1", "toc 2", "toc 3", "toc 4", "toc 5", "toc 6", "toc 7",
                  "toc 8", "toc 9"]
pseudo_code_styles = ["Normal", "Normal_P1", "Normal_P2", "List Paragraph", "Indent 1", "Indent 2", "Req_Text", "Req_Text_L1",
                      "Req_Text_L2", "Req_Text_L3", "Req_Text_P1", "Req_Text_P2", "Req_Text_P3",  "Hyperlink"]
exclude_styles = ["Tableau", "Req_ID"]

open_loop_keywords = ["IF", "CASE", "FOR",
                      "WHILE", "ELSE", "ELSIF", "WHEN"]
close_loop_keyword = ["ENDIF", "ENDCASE", "ENDFOR",
                      "ENDWHILE", "ELSE", "ELSIF", "WHEN"]

file_name_bad_chars = ['/', '\\', '"', ';', ':', '!', "*"]
hierarchy_diagrams_list = []

current_pspec_number = r'P-Spec x.y..'
current_pspec_name = r'name_not_found'
current_diagram_number = r'Diagram x.y..'
current_diagram_name = r'name_not_found'
open_loop_count = 0
space_str = ""
alignement = "    "
is_it_procedure = False
para_text = ""
line = ""


def GetTag(element):
    return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))
# Until hyperlinks functionality is implemented in python-docx, this is the workaround to redefine the 'text' property of the docx.text.paragraph.Paragraph class such that it includes hyperlinks.


def GetParagraphText(paragraph):

    text = ''
    runCount = 0
    for child in paragraph._p:
        tag = GetTag(child)
        if tag == "w:r":
            text += paragraph.runs[runCount].text
            runCount += 1
        if tag == "w:hyperlink":
            for subChild in child:
                if GetTag(subChild) == "w:r":
                    text += subChild.text
    return text


def table_print(block, dir_name):
    table = block
    for row in table.rows:
        column = 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                write_to_file(paragraph, column, dir_name)
                # print(paragraph.text, '  ', end='')
                """
                print('table : ', paragraph.style.name)
                """
                # print(paragraph.text, '\t\n\t', end='')
                # y.write(paragraph.text)
                #y.write('  ')
        # print("\n")
        # y.write("\n")


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def update_diagram_hierarchy_list(current_diagram_number, current_diagram_name):
    no_of_digits_curr_diagram = sum(c.isdigit()
                                    for c in current_diagram_number)

    if no_of_digits_curr_diagram > 0:
        if len(hierarchy_diagrams_list) >= no_of_digits_curr_diagram:
            hierarchy_diagrams_list[no_of_digits_curr_diagram - 1] = (
                current_diagram_number + ' ' + current_diagram_name)

        else:
            hierarchy_diagrams_list.insert(no_of_digits_curr_diagram - 1,
                                           (current_diagram_number + ' ' + current_diagram_name))

        # hierarchy list cleanup of previous length
        while (no_of_digits_curr_diagram < len(hierarchy_diagrams_list)):
            index = (len(hierarchy_diagrams_list) - no_of_digits_curr_diagram)
            hierarchy_diagrams_list.pop(len(hierarchy_diagrams_list) - index)


def update_diagram_hierarchy_list_for_pspec(current_pspec_number, current_diagram_number):
    if (len(hierarchy_diagrams_list) > 1):
        no_of_digits_curr_pspec = sum(c.isdigit()
                                      for c in current_pspec_number)
        no_of_digits_curr_diagram = sum(
            c.isdigit() for c in hierarchy_diagrams_list[len(hierarchy_diagrams_list) - 1])
        # hierarchy list cleanup of previous length
        if (no_of_digits_curr_diagram >= no_of_digits_curr_pspec):
            no_of_pops = (no_of_digits_curr_diagram -
                          no_of_digits_curr_pspec) + 1
            while (no_of_pops):
                no_of_pops = no_of_pops - 1
                if (len(hierarchy_diagrams_list) > 1):
                    hierarchy_diagrams_list.pop(
                        len(hierarchy_diagrams_list) - 1)


def write_to_file(para, column, dir_name):
    global current_pspec_number
    global current_pspec_name
    global current_diagram_number
    global current_diagram_name
    global open_loop_count
    global space_str
    global alignement
    global is_it_procedure
    global para_text
    global line
    para_text = GetParagraphText(para)
    if para.style.name in heading_styles:
        str_list = para_text.split(';')
        if is_diagram(str_list[0].strip()):
            # holds heading before ";" string
            current_diagram_number = str_list[0].strip()
            current_pspec_number = current_diagram_number  # reset on every diagram detection
            if (len(str_list) > 1):
                current_diagram_name = str_list[1].strip()
                for ch in file_name_bad_chars:
                    current_diagram_name = current_diagram_name.replace(
                        ch, '')
                current_pspec_name = 'desciption'  # reset on every diagram detection
            # print("\t\t\t", current_heading_diagram, diagram)
            update_diagram_hierarchy_list(
                current_diagram_number, current_diagram_name)
            is_it_procedure = False
            open_loop_count = 0

        if is_pspec(str_list[0].strip()):
            # holds heading before ";" string
            current_pspec_number = str_list[0].strip()
            if (len(str_list) > 1):
                # holds heading after ";" string
                current_pspec_name = str_list[1].strip()
                for ch in file_name_bad_chars:
                    current_pspec_name = current_pspec_name.replace(ch, '')
            # print("\t\t",current_heading_pspec.strip(),"::",current_heading_name.strip())
            update_diagram_hierarchy_list_for_pspec(
                current_pspec_number, current_diagram_number)
            is_it_procedure = False
            open_loop_count = 0

        if is_procedure(str_list[0].strip()):
            is_it_procedure = True
            open_loop_count = 0

    # update dir_path as per current hierarchy
    dir_path = dir_name
    for name in hierarchy_diagrams_list:
        dir_path = os.path.join(dir_path, name)
        # dir_path = dir_path.replace('Diagram', 'D')
        # dir_path = dir_path.replace(' ', '_')

        if not (os.path.exists(dir_path)):
            os.mkdir(dir_path)

    # write_file_name = (current_pspec_number + current_pspec_name).replace('P-Spec', 'P')
    write_file_name = (current_pspec_number + ' ' + current_pspec_name)
    # write_file_name = write_file_name.replace(' ', '_') + ".pseudo"
    write_file_name = write_file_name + ".gfd"

    if is_it_procedure:
        # procedures are at root folder level
        write_file_path = os.path.join(dir_name, write_file_name)
    else:
        write_file_path = os.path.join(dir_path, write_file_name)

    with open(write_file_path, "a+", encoding="utf-8") as out_file:
        if len(para_text) > 0:
            if (para.style.name in pseudo_code_styles):  # pseudo code syntax
                if (para.style.name.strip() == "Req_Text_L1") or (para.style.name.strip() == "Req_Text_P1"):
                    open_loop_count = 1
                elif (para.style.name.strip() == "Req_Text_L2") or (para.style.name.strip() == "Req_Text_P2"):
                    open_loop_count = 2
                elif (para.style.name.strip() == "Req_Text_L3") or (para.style.name.strip() == "Req_Text_P3"):
                    open_loop_count = 3
                elif (para.style.name.strip() == "List Paragraph"):
                    open_loop_count = column
                else:
                    open_loop_count = 0

                space_str = "{}".format(alignement * open_loop_count)

                # in IF statement & conditions were not aligned correctly hence replaced the newline
                line = "{}{}".format(
                    space_str,  para_text.replace("\n", "\n" + space_str))
            else:
                # remove newline in comments
                line = "{}{}{}".format(
                    space_str, "// ", para_text.replace("\n", " "))
            if not (para.style.name in exclude_styles):
                line = line + '\n'
                print(line, file=out_file)


# This procedure extracts all pspec's under matching diagram .
def extract_diagrams_and_pspecs(docx_file, dir_name):

    doc = docx.Document(docx_file)
    # document = Document('test.docx')
# for block in iter_block_items(document):
#     print('found one')
#     print(block.text if isinstance(block, Paragraph) else '<table>')
    print("SRD docx is being parsed : It may take around 2-mins")
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            column = 1
            write_to_file(block, column, dir_name)
        elif isinstance(block, Table):
            table_print(block, dir_name)
        elif isinstance(block, Image):
            print('found image')

    """
    for para in doc.paragraphs:
        # print(para.style.name)
        # for run_inst in para.runs:
        # print(run_inst.text)
        print(GetParagraphText(para))
    """


def is_diagram(first_str):
    if type(first_str) == str:
        str_diagram_space = first_str.split(" ")
        str_diagram_dot = first_str.split(".")
        if ("Diagram".lower() == str_diagram_space[0].lower()) or ("Diagram".lower() == str_diagram_dot[0].lower()):
            return True
        else:
            return False
    else:
        return False


def is_pspec(first_str):
    if type(first_str) == str:
        str_diagram_space = first_str.split(" ")
        str_diagram_dot = first_str.split(".")
        if ("P-Spec".lower() == str_diagram_space[0].lower()) or (
                "Procedure".lower() == str_diagram_space[0].lower()) or \
                ("P-Spec".lower() == str_diagram_dot[0].lower()) or ("Procedure".lower() == str_diagram_dot[0].lower()):
            return True
        else:
            return False
    else:
        return False


def is_procedure(first_str):
    if type(first_str) == str:
        str_diagram_space = first_str.split(" ")
        str_diagram_dot = first_str.split(".")
        if ("Procedure".lower() == str_diagram_space[0].lower()) or ("Procedure".lower() == str_diagram_dot[0].lower()):
            return True
        else:
            return False
    else:
        return False


def delete_dir_and_add_dir(script_dir, dir):
    path = os.path.join(script_dir, dir)
    print(path)
    if os.path.exists(path):
        shutil.rmtree(path)
        print("% s has been removed successfully" % dir)
    time.sleep(1)
    if not os.path.exists(path):
        os.mkdir(path)
        print("% s has been created successfully" % dir)


def check_if_current_diagram_in_specified_list(diagram_number, pspec_name_list):
    to_remove_chars = ['.', " "]
    if type(diagram_number) == str:
        for ch in to_remove_chars:
            diagram_number = diagram_number.replace(ch, '')
        diagram_number = diagram_number.lower()
        # print("diagram_number : ",diagram_number )
        if (diagram_number.strip() in pspec_name_list):
            return True
        else:
            return False
    else:
        return False


def check_if_current_pspec_in_specified_list(pspec_number, pspec_name_list):
    to_remove_chars = ['.', " "]
    if type(pspec_number) == str:
        for ch in to_remove_chars:
            pspec_number = pspec_number.replace(ch, '')
        pspec_number = pspec_number.lower()
        # print( "pspec_number : ",pspec_number)
        if (pspec_number.strip() in pspec_name_list):
            return True
        else:
            return False
    else:
        return False


# This function is for development / debug.  which list all style names in a documents
def list_all_styles_in_docx(docx_file):
    doc = docx.Document(docx_file)
    with open("style_name.txt", "a+", encoding="utf-8") as out_file:
        for para in doc.paragraphs:
            if len(para.style.name) > 0:
                print(para.style.name, file=out_file)


def main(argv):
    # Get the SRD names , where changes are accepted in one and changes are rejected in another for a CR
    no_of_arguments = len(argv)

    script_dir = os.path.dirname(os.path.abspath(
        inspect.getfile(inspect.currentframe())))

    # Extracting the specified pspecs from SRD with accepted changes for a CR
    ct = datetime.datetime.now()
    print("current time:-", ct)

    srd_file_name_accepted = argv[0].strip()
    delete_dir_and_add_dir(script_dir, "accepted")
    srd_file_path = os.path.join(script_dir, srd_file_name_accepted)
    dir_name = os.path.join(script_dir, "accepted")
    extract_diagrams_and_pspecs(srd_file_path, dir_name)

    ct = datetime.datetime.now()
    print("current time:-", ct)

    if (no_of_arguments > 1):
        srd_file_name_rejected = argv[1].strip()
        # Extracting the specified pspecs from SRD with rejected changes for a CR
        delete_dir_and_add_dir(script_dir, "rejected")
        srd_file_path = os.path.join(script_dir, srd_file_name_rejected)
        dir_name = os.path.join(script_dir, "rejected")
        extract_diagrams_and_pspecs(srd_file_path, dir_name)

        ct = datetime.datetime.now()
        print("current time:-", ct)


# End of main

if __name__ == "__main__":
    if len(sys.argv) == 2 or len(sys.argv) == 3:
        main(sys.argv[1:])
    else:
        print("usage is : \n"
              "python extract_all_pspecs.py Rs_T_A400392-xx_accepted.docx \n" "OR \n"
              "python extract_all_pspecs.py Rs_T_A400392-xx_accepted.docx Rs_T_A400392-xx_rejected.docx")
